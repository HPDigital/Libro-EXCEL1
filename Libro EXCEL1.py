"""
Libro EXCEL1
"""

#!/usr/bin/env python
# coding: utf-8

# In[6]:


from docx import Document
import os
from dotenv import load_dotenv
from openai import OpenAI

# Cargar variables de entorno
load_dotenv()

# Verificar si la API key se carga correctamente
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("API Key no encontrada. Asegúrate de que 'OPENAI_API_KEY' está definida en el archivo .env")

# Inicializar el cliente de OpenAI
client = OpenAI(api_key=api_key)

# Temas a procesar
TEMAS = [
    ("Filtros", "PESTAÑA INSERTAR"),
    ("Herramientas de dibujo", "PESTAÑA DIBUJAR"),
    ("Convertir", "PESTAÑA DIBUJAR"),
    ("Reproducción", "PESTAÑA DIBUJAR"),
    ("Ayuda", "PESTAÑA DIBUJAR")
    ]
    # ("Portapapeles", "PESTAÑA INICIO"), 
    # ("Fuente", "PESTAÑA INICIO"),
    # ("Alineación", "PESTAÑA INICIO"),
    # ("Numero", "PESTAÑA INICIO"),
    # ("Estilos", "PESTAÑA INICIO"),
    # ("Celdas", "PESTAÑA INICIO"),
    # ("Edición", "PESTAÑA INICIO"),
    # ("Complementos", "PESTAÑA INICIO"),
    # ("Tablas", "PESTAÑA INSERTAR"),
    # ("Ilustraciones", "PESTAÑA INSERTAR"),
    # ("Controles", "PESTAÑA INSERTAR"),
    # ("Gráficos", "PESTAÑA INSERTAR"),
    # ("Filtros", "PESTAÑA INSERTAR"),
    # ("Vínculos", "PESTAÑA INSERTAR"),
    # ("Comentarios", "PESTAÑA INSERTAR"),
    # ("Texto", "PESTAÑA INSERTAR"),
    # ("Símbolos", "PESTAÑA INSERTAR"),
    # ("Temas", "PESTAÑA DISPOSICIÓN DE PAGINA"),
    # ("Configurar página", "PESTAÑA DISPOSICIÓN DE PAGINA"),
    # ("Ajustar área de impresión", "PESTAÑA DISPOSICIÓN DE PAGINA"),
    # ("Hoja", "PESTAÑA DISPOSICIÓN DE PAGINA"),
    # ("Organizar", "PESTAÑA DISPOSICIÓN DE PAGINA"),
    # ("Biblioteca de funciones", "PESTAÑA FÓRMULAS"),
    # ("Python Versión preliminar", "PESTAÑA FÓRMULAS"),
    # ("Nombres definidos", "PESTAÑA FÓRMULAS"),
    # ("Auditoría de fórmulas", "PESTAÑA FÓRMULAS"),
    # ("Cálculo", "PESTAÑA FÓRMULAS"),
    # ("Obtener y transformar datos", "PESTAÑA DATOS"),
    # ("Consultas y conexiones", "PESTAÑA DATOS"),
    # ("Tipos de datos", "PESTAÑA DATOS"),
    # ("Ordenar y Filtrar", "PESTAÑA DATOS"),
    # ("Herramientas de datos", "PESTAÑA DATOS"),
    # ("Previsión", "PESTAÑA DATOS"),
    # ("Esquema", "PESTAÑA DATOS"),
    # ("Revisión", "PESTAÑA REVISAR"),
    # ("Rendimiento", "PESTAÑA REVISAR"),
    # ("Accesibilidad", "PESTAÑA REVISAR"),
    # ("Datos", "PESTAÑA REVISAR"),
    # ("Idioma", "PESTAÑA REVISAR"),
    # ("Comentarios", "PESTAÑA REVISAR"),
    # ("Notas", "PESTAÑA REVISAR"),
    # ("Proteger", "PESTAÑA REVISAR"),
    # ("Vista de hoja", "PESTAÑA VISTA"),
    # ("Vistas de estilo", "PESTAÑA VISTA"),
    # ("Mostrar", "PESTAÑA VISTA"),
    # ("Zoom", "PESTAÑA VISTA"),
    # ("Ventana", "PESTAÑA VISTA"),
    # ("Código", "PESTAÑA PROGRAMADOR"),
    # ("Complementos", "PESTAÑA PROGRAMADOR"),
    # ("Controles", "PESTAÑA PROGRAMADOR"),
    # ("XML", "PESTAÑA PROGRAMADOR"),
    # ("Ayuda", "PESTAÑA AYUDA"),

    # ("Modelo de datos", "PESTAÑA POWER PIVOT"),
    # ("Cálculos", "PESTAÑA POWER PIVOT"),
    # ("Tablas", "PESTAÑA POWER PIVOT"),
    # ("Relaciones", "PESTAÑA POWER PIVOT"),
    # ("Inicio", "PESTAÑA ARCHIVO"),
    # ("Nuevo", "PESTAÑA ARCHIVO"),
    # ("Abrir", "PESTAÑA ARCHIVO"),
    # ("Obtener complementos", "PESTAÑA ARCHIVO"),
    # ("Información", "PESTAÑA ARCHIVO"),
    # ("Guardar", "PESTAÑA ARCHIVO"),
    # ("Guardar como", "PESTAÑA ARCHIVO"),
    # ("Imprimir", "PESTAÑA ARCHIVO"),
    # ("Compartir", "PESTAÑA ARCHIVO"),
    # ("Exportar", "PESTAÑA ARCHIVO"),
    # ("Cerrar", "PESTAÑA ARCHIVO"),
    # ("Cuenta", "PESTAÑA ARCHIVO"),
    # ("Opciones", "PESTAÑA ARCHIVO")
    # ]


# Plantilla del prompt
PROMPT_TEMPLATE = (
    "Actúa como un experto en Excel y un profesor con experiencia en enseñar Excel a estudiantes de todos los niveles. "
    "Proporciona una explicación exhaustiva, clara y detallada sobre cada una de las funciones disponibles en el Grupo de '{grupo}' "
    "de la '{pestaña}' en Excel. Asegúrate de describir cada función en profundidad, incluir su propósito, cómo se utiliza, los pasos "
    "detallados para aplicarla y dar ejemplos prácticos que muestren situaciones específicas en las que sería útil. No omitas ninguna función del Grupo "
    "de '{grupo}' e incluye información sobre cualquier opción o característica avanzada relacionada con estas herramientas. Además, explica cómo combinarlas "
    "para optimizar flujos de trabajo en Excel y menciona los atajos de teclado relevantes para cada función. Incluye consejos para evitar errores comunes "
    "al usarlas y destaca su importancia en la gestión eficiente de datos en Excel."
    "el tono de la respuesta debe ser muy academico, y no debe haber en la respuesta formulas de cortesia"
)

responses = []

for grupo, pestaña in TEMAS:
    prompt = PROMPT_TEMPLATE.format(grupo=grupo, pestaña=pestaña)
    try:
        # Ajustar el formato del llamado según la API actual
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Eres un experto en Excel y un profesor experimentado."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000
        )
        responses.append(response.choices[0].message.content)
    except Exception as e:
        print(f"Error al procesar el tema '{grupo} - {pestaña}': {e}")

# Crear un documento Word y guardar las respuestas
doc = Document()
doc.add_heading("Explicaciones de Temas de Excel", level=1)

for response in responses:
    doc.add_paragraph(response)
    doc.add_paragraph()  # Línea en blanco entre explicaciones

output_path = r"C:\Users\HP\Desktop\LIBROS PERSO\EXCEL INTERFAZ GRAFICA\temas_excel_explicaciones.docx"
doc.save(output_path)

print(f"Se ha generado el archivo en '{output_path}' con las explicaciones de los temas.")


# In[ ]:




